# app/helper/document_extractor.py
import io
import os
import json
import csv
import docx
import tempfile

from pathlib import Path
from typing import List
from fastapi import HTTPException

from llama_index.core import Document
from llama_index.readers.docling import DoclingReader

from app.core.enums.file import FileExtension, EncodingType, ParsingConstants
from app.core.enums.http import ErrorMessage, HttpStatus
from app.core.logger import get_logger

logger = get_logger("extractor")

def extract_documents(file_bytes: bytes, filename: str, reader_map: dict, arg_map: dict) -> List[Document]:
    """
    Try DoclingReader first. If it fails, fallback to LlamaIndex-specific readers.
    If all fail, fallback to manual extractor.
    """
    ext = os.path.splitext(filename.lower())[1]
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        logger.info(f"[Extractor] DoclingReader extractor {filename}")
        docling_reader = DoclingReader()
        docs = docling_reader.load_data(file_path=tmp_path)
        for doc in docs:
            doc.metadata.update({
                "file_name": filename,
                "file_type": ext,
                "processing_method": "docling",
                "reader_type": "DoclingReader"
            })
        os.remove(tmp_path)
        return docs
    except Exception as e:
        logger.warning(f"[Extractor] DoclingReader failed for {filename}: {e}")

    logger.info(f"[Extractor] LlamaIndex Default Reader extractor {filename}")
    if ext in reader_map:
        try:
            reader = reader_map[ext]
            arg_name = arg_map.get(ext, "file")
            docs = reader.load_data(**{arg_name: Path(tmp_path)})

            for doc in docs:
                doc.metadata.update({
                    "file_name": filename,
                    "file_type": ext,
                    "processing_method": "llama_index_reader",
                    "reader_type": type(reader).__name__
                })
            os.remove(tmp_path)
            return docs
        except Exception as e:
            logger.warning(f"[Extractor] LlamaIndex reader failed for {filename}: {e}")

        logger.info(f"[Extractor] Fallback to manual extraction for {filename}")
        text = _extract_text_manual(file_bytes, filename)
        return [Document(text=text, metadata={"file_name": filename})]


def _extract_text_manual(file_bytes: bytes, filename: str) -> str:
    """
    Dispatch manual extraction function based on file extension.
    """
    ext = os.path.splitext(filename.lower())[1]
    match ext:
        case FileExtension.TXT:
            return _extract_text_from_txt(file_bytes)
        case FileExtension.DOCX:
            return _extract_text_from_docx(file_bytes)
        case FileExtension.CSV:
            return _extract_text_from_csv(file_bytes)
        case FileExtension.JSON:
            return _extract_text_from_json(file_bytes)
        case _:
            raise HTTPException(
                status_code=HttpStatus.UNPROCESSABLE_ENTITY.value,
                detail=f"{ErrorMessage.UNSUPPORTED_FILE_TYPE}: {ext}"
            )

def _extract_text_from_txt(binary_txt: bytes) -> str:
    """
    Try decoding a text file using multiple encodings.
    """
    for enc in [
        EncodingType.UTF8,
        EncodingType.UTF8_SIG,
        EncodingType.LATIN1,
        EncodingType.CP1252
    ]:
        try:
            return binary_txt.decode(enc)
        except UnicodeDecodeError:
            continue
    raise HTTPException(
        status_code=HttpStatus.UNPROCESSABLE_ENTITY.value,
        detail=ErrorMessage.UNABLE_TO_DECODE_TEXT
    )


def _extract_text_from_docx(binary_docx: bytes) -> str:
    """
    Extract all text from a DOCX file including table content.
    """
    doc = docx.Document(io.BytesIO(binary_docx))
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(ParsingConstants.COLUMN_SEPARATOR.join(row_text))

    return '\n'.join(text_parts)


def _extract_text_from_csv(binary_csv: bytes) -> str:
    """
    Extract structured text from a CSV file.
    """
    text_content = binary_csv.decode(EncodingType.UTF8_SIG)
    csv_reader = csv.reader(io.StringIO(text_content))
    text_parts = []
    headers = None

    for i, row in enumerate(csv_reader):
        if i == 0:
            headers = row
            text_parts.append(ParsingConstants.HEADERS_PREFIX + ParsingConstants.COLUMN_SEPARATOR.join(headers))
        else:
            if headers and len(row) == len(headers):
                row_dict = dict(zip(headers, row))
                row_text = ParsingConstants.COLUMN_SEPARATOR.join([
                    f"{k}{ParsingConstants.KEY_VALUE_SEPARATOR}{v}" for k, v in row_dict.items() if v.strip()
                ])
            else:
                row_text = ParsingConstants.COLUMN_SEPARATOR.join([cell for cell in row if cell.strip()])
            text_parts.append(row_text)

    return '\n'.join(text_parts)


def _extract_text_from_json(binary_json: bytes) -> str:
    """
    Extract structured text from a nested JSON file.
    """
    text_content = binary_json.decode(EncodingType.UTF8_SIG)
    data = json.loads(text_content)

    def extract_text(obj, prefix="") -> List[str]:
        parts = []
        if isinstance(obj, dict):
            for k, v in obj.items():
                current = f"{prefix}.{k}" if prefix else k
                parts.extend(extract_text(v, current))
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                current = f"{prefix}[{i}]" if prefix else f"item[{i}]"
                parts.extend(extract_text(item, current))
        else:
            parts.append(f"{prefix}: {str(obj)}")
        return parts

    return '\n'.join(extract_text(data))
